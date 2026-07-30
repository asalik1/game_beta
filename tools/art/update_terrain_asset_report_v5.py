"""Append the TERRAIN_ART_FIX_TASK all-tier pass to the terrain art report."""

from __future__ import annotations

from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "Crownless_Terrain_Asset_Changes_2026-07-29_v4.docx"
OUTPUT = ROOT / "docs" / "Crownless_Terrain_Asset_Changes_2026-07-30_v5.docx"
QA_SOURCE = ROOT / "tmp" / "terrain_art_fix_qa" / "all_tiers.png"
QA_DIR = ROOT / "tmp" / "terrain_report_v5_qa"

ASSETS = [
    ("1", "cottage_a", "384x300", "Distinct thatch cottage; masonry base and readable entrance"),
    ("1", "cottage_a2", "384x320", "Independent mossed cottage silhouette, not a recolor"),
    ("1", "cottage_b", "384x260", "Broad stone dwelling with authored roof and facade depth"),
    ("1", "stall", "320x252", "Grounded timber market stall with stocked counter"),
    ("1", "rock3", "256x320", "Painterly stone cluster; magenta fringe removed"),
    ("1", "crypt", "256x300", "Complete walk-in mausoleum with dimensional masonry"),
    ("1", "signpost", "112x192", "Readable timber wayfinder with stone-set footing"),
    ("2", "keep_arch", "320x240", "Ruined masonry arch with coherent depth and rubble"),
    ("2", "camp_workbench", "288x240", "Tool-rich field bench without heavy cartoon outlines"),
    ("2", "cook_grill", "256x224 + 4f", "Complete four-frame grill; coals and food animate"),
    ("2", "camp_bonfire", "192x128 + 4f", "Complete four-frame fire; stone ring stays grounded"),
    ("2", "pillar", "160x256", "Weathered carved pillar with full base and capital"),
    ("3", "banner_blue", "96x192 + 4f", "Complete four-frame blue cloth billow"),
    ("3", "banner_green", "96x192 + 4f", "Complete four-frame green cloth billow"),
    ("3", "banner_red", "96x192 + 4f", "Complete four-frame red cloth billow"),
    ("3", "hideout_poster", "96x144", "Aged, pinned faction notice with readable silhouette"),
    ("3", "hideout_table", "256x224", "Dimensional planning table with map and candle"),
    ("3", "amphora", "112x192", "Detailed terracotta vessel with strong handles"),
    ("3", "station_alchemy_t3", "320x256 + 4f", "Master station; contained liquid and gauges animate"),
    ("3", "station_anvil_t3", "320x288", "Complete master smithing station and tool wall"),
]


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _split_contact_sheet() -> tuple[Path, Path]:
    QA_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.open(QA_SOURCE).convert("RGB")
    midpoint = image.height // 2
    top = QA_DIR / "all_tiers_top.png"
    bottom = QA_DIR / "all_tiers_bottom.png"
    image.crop((0, 0, image.width, midpoint)).save(top, quality=94)
    image.crop((0, midpoint, image.width, image.height)).save(bottom, quality=94)
    return top, bottom


def _add_contact_page(doc: Document, image_path: Path, caption_text: str) -> None:
    doc.add_page_break()
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(image_path), width=Inches(6.35))
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document(SOURCE)
    top, bottom = _split_contact_sheet()

    doc.add_page_break()
    label = doc.add_paragraph()
    run = label.add_run("TERRAIN ART FIX — ALL TIERS")
    run.bold = True
    doc.add_heading("Twenty direct scenery replacements", level=1)
    doc.add_paragraph(
        "The complete TERRAIN_ART_FIX_TASK register is now replaced in production, "
        "covering all seven Tier 1 assets, all five Tier 2 assets, and all eight "
        "Tier 3 assets. These are direct high-resolution replacements for the "
        "existing terrain objects, not detached decoration layered over the old art."
    )

    summary = doc.add_table(rows=5, cols=2)
    summary.style = doc.tables[0].style
    summary_values = [
        ("Measure", "Delivered"),
        ("Tier coverage", "Tier 1: 7 / 7   |   Tier 2: 5 / 5   |   Tier 3: 8 / 8"),
        ("Production replacements", "20 static assets"),
        ("Full-object motion", "6 four-frame strips; no child effect overlays"),
        ("Desktop / mobile", "Byte-identical source PNGs and mirrored runtime logic"),
    ]
    for row, values in zip(summary.rows, summary_values):
        for cell, value in zip(row.cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _repeat_header(summary.rows[0])

    doc.add_heading("Asset register", level=2)
    register = doc.add_table(rows=1, cols=4)
    register.style = doc.tables[0].style
    headers = ("Tier", "Asset", "Production canvas", "Direct improvement")
    for cell, value in zip(register.rows[0].cells, headers):
        cell.text = value
        _shade(cell, "D9E2F3")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _repeat_header(register.rows[0])
    widths = (Inches(0.45), Inches(1.45), Inches(1.25), Inches(3.25))
    for tier, name, canvas, improvement in ASSETS:
        row = register.add_row()
        values = (tier, name, canvas, improvement)
        for cell, value, width in zip(row.cells, values, widths):
            cell.text = value
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    _add_contact_page(
        doc,
        top,
        "Production contact sheet, assets 1–10. Checkerboard denotes transparency.",
    )
    _add_contact_page(
        doc,
        bottom,
        "Production contact sheet, assets 11–20. Animated entries show all four complete frames.",
    )

    doc.add_page_break()
    doc.add_heading("Runtime integration and validation", level=1)
    for text in (
        "Explicit render widths preserve sensible world scale despite the new larger source canvases.",
        "Authored collider radii now cover the logical footprints of the crypt, arch, camp objects, pillar, table, amphora, and tier-three stations.",
        "GameWorld builds animated scenery as one complete AnimatedSprite2D. The six new strips never use a pasted child animation.",
        "Ruined-gate and market banners no longer combine their full-object strips with a second wind deformation.",
        "The complete source/edit/build path is archived under art_src/terrain_art_fix_2026-07-30/ and tools/art/build_terrain_art_fix.py.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Generation contract", level=2)
    doc.add_paragraph(
        "Built-in image generation was used in reference-image edit mode, one "
        "asset at a time. Prompts used the original prop for identity and the "
        "project's capital scenery, volcanic rock, statues, and ruin pillar for "
        "quality direction: painterly chunky pixel clusters, clear top-down "
        "three-quarter silhouettes, restrained dark-fantasy materials, flat "
        "magenta chroma background, no floor, no cast shadow, no text, and no "
        "photorealistic or smooth 3D finish. Animated prompts required four "
        "equal horizontal complete-object frames with a fixed canvas and rigid "
        "structure; only fire, cloth, contained liquid, or gauges could change."
    )

    checks = doc.add_table(rows=7, cols=3)
    checks.style = doc.tables[0].style
    check_values = [
        ("Check", "Result", "Evidence"),
        ("Task coverage", "PASS", "20 / 20 assets across all three tiers"),
        ("Art verifier", "PASS", "20 static sources and 6 strips imported cleanly"),
        ("Transparency", "PASS", "Binary alpha; no surviving magenta fringe"),
        ("Frame construction", "PASS", "Four complete frames; static-size match"),
        ("Desktop quick suite", "PASS", "93 scripts; scenery regression checks green"),
        ("Mobile quick suite", "PASS", "93 scripts; mirrored checks green"),
    ]
    for row, values in zip(checks.rows, check_values):
        for cell, value in zip(row.cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _repeat_header(checks.rows[0])
    doc.add_paragraph("Report revision: v5 — 30 July 2026")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
