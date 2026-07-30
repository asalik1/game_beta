"""Append the full-object animation correction to the terrain DOCX report."""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "Crownless_Terrain_Asset_Changes_2026-07-28_v3.docx"
OUTPUT = ROOT / "docs" / "Crownless_Terrain_Asset_Changes_2026-07-29_v4.docx"
SPRITES = ROOT / "game" / "assets" / "sprites"
QA_IMAGE = ROOT / "tmp" / "terrain_prop_anims_report_contact.png"

ASSETS = [
    ("garden_fountain", "Water spout, streams and basin ripples"),
    ("spore_vent", "Inner violet pulse and spores"),
    ("void_rift", "Contained void field and energy veins"),
    ("capital_portal_depths", "Portal field and restrained motes"),
    ("storm_conductor", "Attached cyan electrical arcs"),
    ("magma_furnace", "Coals and fire inside the furnace mouth"),
    ("keep_brazier", "Integrated brazier flame and coals"),
    ("forge_cauldron", "Integrated crucible fire"),
    ("forge_brazier", "Integrated small-brazier fire"),
    ("camp_furnace", "Integrated field-furnace fire"),
    ("station_furnace_t1", "Tier-one furnace fire"),
    ("station_furnace_t2", "Tier-two furnace fire"),
    ("station_furnace_t3", "Tier-three furnace fire"),
    ("sewer_outfall", "Pipe-connected wastewater and puddle"),
]


def _replace_text(doc: Document) -> None:
    replacements = {
        "48 generated sprites  •  14 live terrain passes  •  20 new placeholders":
            "48 generated sprites  •  14 live terrain passes  •  20 new dev-preview terrains",
        "Placeholder register": "Dev-preview terrain register",
        "DEV-PANEL PLACEHOLDERS": "DEV-PANEL TERRAIN PREVIEWS",
        "The 20 new terrain placeholders, kept separate from the 16 older placeholder profiles.":
            "The 20 new non-placeholder terrain previews, kept separate from the 16 older placeholder profiles.",
        "28 July 2026": "28-29 July 2026",
    }
    for paragraph in doc.paragraphs:
        original = paragraph.text
        if original in replacements:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = replacements[original]

        text = paragraph.text
        if text.startswith("The 20 new terrain IDs are tagged placeholder: true"):
            paragraph.text = (
                "The 20 new terrain IDs are marked preview_isolated: true and "
                "remain visible in the dev panel / Future → Terrains gallery. "
                "They are not placeholder entries, none is referenced by "
                "Story.CHAPTERS, and no Chapter 3–7 zone was reassigned."
            )
        elif text == "All new terrain concepts remain placeholder: true.":
            paragraph.text = (
                "All 20 new terrain concepts are non-placeholder, isolated "
                "dev-preview entries."
            )
        elif text.startswith("Twenty terrain IDs remain placeholder: true"):
            paragraph.text = (
                "Twenty terrain IDs are non-placeholder dev-preview entries "
                "exposed for terrain repainting; no Chapters 3–7 zone was reassigned."
            )
        elif text.startswith("Terrains.PROP_MOTION now describes local active-element overlays"):
            paragraph.text = (
                "Superseded on 29 July: PROP_MOTION and its child-overlay runtime "
                "were removed. Every affected prop now loads one complete "
                "<name>_anim.png strip through the shared prop renderer."
            )
        elif text.startswith("Furnace, cauldron and brazier families receive animated fire"):
            paragraph.text = (
                "Furnace, cauldron and brazier families now contain their fire "
                "inside four complete, footprint-anchored object frames."
            )
        elif text.startswith("town_fountain, garden_fountain and holy_sanctum share"):
            paragraph.text = (
                "town_fountain, garden_fountain and holy_sanctum now share the "
                "same complete animated fountain base; no fountain_flow decal is used."
            )
        elif text == "All twenty remain visible through the development terrain panel.":
            paragraph.text = (
                "All twenty remain visible through the development terrain panel "
                "as non-placeholder, isolated previews."
            )
        elif text.startswith("Each placeholder now resolves"):
            paragraph.text = text.replace(
                "Each placeholder",
                "Each dev-preview terrain",
                1,
            )
        elif text.startswith("Placeholder previews ignore"):
            paragraph.text = text.replace(
                "Placeholder previews",
                "Dev previews",
                1,
            )

    cell_replacements = {
        "20 dev-panel placeholders": "20 non-placeholder dev previews",
        "20 placeholder definitions;": "20 dev-preview definitions;",
        "all IDs remain placeholder: true.": "all IDs are non-placeholder dev previews.",
        "placeholder status": "preview-isolation status",
        "Placeholder audit": "Dev-preview audit",
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for old, new in cell_replacements.items():
                        text = text.replace(old, new)
                    if text != paragraph.text:
                        paragraph.text = text


def _remove_forced_break_before(doc: Document, target: str) -> None:
    """Remove an empty page-break paragraph immediately before target."""
    paragraphs = doc.paragraphs
    target_index = next(
        (index for index, paragraph in enumerate(paragraphs) if paragraph.text == target),
        None,
    )
    if target_index is None:
        raise RuntimeError(f"Could not find report section: {target}")
    for index in range(target_index - 1, -1, -1):
        paragraph = paragraphs[index]
        if paragraph.text:
            break
        if paragraph._p.xpath(".//w:br"):
            parent = paragraph._element.getparent()
            parent.remove(paragraph._element)
            return
    raise RuntimeError(f"Could not find forced break before report section: {target}")


def _contact_sheet() -> None:
    width, height = 1500, 1500
    columns, rows = 2, 7
    cell_w, cell_h = width // columns, height // rows
    canvas = Image.new("RGBA", (width, height), (24, 27, 31, 255))
    draw = ImageDraw.Draw(canvas)
    for index, (name, _) in enumerate(ASSETS):
        col, row = index % columns, index // columns
        left, top = col * cell_w, row * cell_h
        draw.rounded_rectangle(
            (left + 12, top + 10, left + cell_w - 12, top + cell_h - 10),
            radius=18,
            fill=(33, 37, 43, 255),
            outline=(83, 91, 101, 255),
            width=2,
        )
        draw.text((left + 28, top + 23), name, fill=(232, 235, 238, 255))
        strip = Image.open(SPRITES / f"{name}_anim.png").convert("RGBA")
        max_w, max_h = cell_w - 52, cell_h - 66
        scale = min(max_w / strip.width, max_h / strip.height)
        scaled = strip.resize(
            (max(1, round(strip.width * scale)), max(1, round(strip.height * scale))),
            Image.Resampling.NEAREST,
        )
        canvas.alpha_composite(
            scaled,
            (
                left + (cell_w - scaled.width) // 2,
                top + 54 + (max_h - scaled.height) // 2,
            ),
        )
    QA_IMAGE.parent.mkdir(parents=True, exist_ok=True)
    canvas.convert("RGB").save(QA_IMAGE, quality=92)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _set_repeat_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def main() -> None:
    doc = Document(SOURCE)
    _replace_text(doc)
    for section in (
        "Authored prop variation families",
        "IMPLEMENTATION",
        "SURFACE REGISTER",
    ):
        _remove_forced_break_before(doc, section)
    _contact_sheet()

    doc.add_page_break()
    label = doc.add_paragraph()
    label_run = label.add_run("FULL-OBJECT ANIMATION CORRECTION")
    label_run.bold = True
    doc.add_heading("Integrated motion assets", level=1)
    doc.add_paragraph(
        "The earlier kinetic-prop addendum solved missing motion but still "
        "composited separate water, flame, void, spore and lightning sprites "
        "over static bodies. That architecture caused the fountain spout to "
        "shift and made active elements look pasted on. This correction "
        "regenerates every affected asset as four complete object frames."
    )

    summary = doc.add_table(rows=4, cols=2)
    summary.style = doc.tables[0].style
    values = [
        ("Measure", "Delivered"),
        ("Complete regenerated props", "14 static sources + 14 four-frame strips"),
        ("Independent motion decals removed", "PROP_MOTION, fountain_flow and sewer_flow paths"),
        ("Frame anchoring", "Identical baseline; centre variance at most 2 px"),
    ]
    for row, values_row in zip(summary.rows, values):
        for cell, value in zip(row.cells, values_row):
            cell.text = value
    _set_repeat_header(summary.rows[0])

    doc.add_paragraph()
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(QA_IMAGE), width=Inches(6.45))
    caption = doc.add_paragraph(
        "Production contact sheet: every cell shows all four complete frames, "
        "not a detached effect layer."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Regenerated asset register", level=2)
    register = doc.add_table(rows=1, cols=3)
    register.style = doc.tables[0].style
    for cell, value in zip(
        register.rows[0].cells,
        ["Asset", "Integrated motion", "Installation"],
    ):
        cell.text = value
    _set_repeat_header(register.rows[0])
    for name, motion in ASSETS:
        row = register.add_row()
        row.cells[0].text = f"{name}.png"
        row.cells[1].text = motion
        row.cells[2].text = f"{name}_anim.png (4 complete frames)"

    doc.add_heading("Runtime and physical integration", level=2)
    _bullet(
        doc,
        "GameWorld._prop_visual now chooses either one full AnimatedSprite2D "
        "or one static Sprite2D; it never adds a child motion overlay.",
    )
    _bullet(
        doc,
        "Fountain water is painted into the complete fountain strip. The "
        "town, garden and holy variants retain their two-lobe basin blockers.",
    )
    _bullet(
        doc,
        "The sewer outfall now uses a dedicated full pipe + wastewater asset "
        "and a wider authored footprint instead of sewer_pipe + sewer_flow.",
    )
    _bullet(
        doc,
        "Keep and magma landmark lighting is preserved through light-only "
        "sockets, without restoring flame or smoke sprites.",
    )
    _bullet(
        doc,
        "Tiny legacy crafting furnaces received higher-resolution canvases "
        "while explicit render widths preserve their previous world scale.",
    )
    _bullet(
        doc,
        "The reproducible build path is tools/art/build_terrain_prop_anims.py; "
        "original keyed sources and transparent masters are archived under "
        "art_src/terrain_prop_anims_2026-07-29/.",
    )

    doc.add_heading("Generation prompt and validation", level=2)
    doc.add_paragraph(
        "Built-in image generation was used in image-edit mode with each "
        "existing in-repo prop as the visual reference. Every prompt required "
        "exactly four equal horizontal frames, the complete prop in every "
        "frame, one stable baseline and position, rigid stone/metal/wood, only "
        "the physically active material changing, and chunky limited-palette "
        "dark-fantasy pixel art on a flat chroma background."
    )
    checks = doc.add_table(rows=6, cols=3)
    checks.style = doc.tables[0].style
    check_values = [
        ("Check", "Result", "Evidence"),
        ("Godot import", "PASS", "28 desktop + 28 mobile sprite files imported"),
        ("Desktop compile / quick suite", "PASS", "93 scripts; full-object seam regressions green"),
        ("Mobile compile / quick suite", "PASS", "93 scripts; same asset seam tests green"),
        ("Geometry regression", "PASS", "4 frames; matching static size; anchored baseline/centre"),
        ("Overlay regression", "PASS", "No nested AnimatedSprite2D on all 14 props"),
    ]
    for row, values_row in zip(checks.rows, check_values):
        for cell, value in zip(row.cells, values_row):
            cell.text = value
    _set_repeat_header(checks.rows[0])
    doc.add_paragraph("Report revision: v4 — 29 July 2026")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
